#!/usr/bin/env python3
"""Convert Crescendo research markdown to Word document"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Crescendo.ai - Enterprise Account Executive', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Interview Preparation Research')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].bold = True

date_para = doc.add_paragraph('Generated: January 7, 2026')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].italic = True

doc.add_paragraph()

# Company Overview
doc.add_heading('Company Overview', 1)

doc.add_heading('Mission & Vision', 2)
p = doc.add_paragraph('Crescendo.ai developed the world\'s first ')
run = p.add_run('augmented-AI customer experience (CX) platform')
run.bold = True
p.add_run(' that combines the power of AI with human domain expertise to deliver outcome-driven results for fast-moving, digitally-native enterprises.')

doc.add_heading('Core Value Proposition', 2)
doc.add_paragraph('AI + Superhumans: Working together to resolve what bots can\'t and BPOs won\'t', style='List Bullet')
doc.add_paragraph('Pay-for-Outcomes Model: Total Outcome Guarantee for quality, speed, and customer satisfaction', style='List Bullet')
doc.add_paragraph('Fully Automated Tier 1: Instant, multilingual support available 24/7', style='List Bullet')

doc.add_heading('Products & Services', 2)
doc.add_paragraph('The Crescendo Platform:', style='List Bullet')
doc.add_paragraph('Directs foundational AI models to conduct sales or customer support conversations', style='List Bullet 2')
doc.add_paragraph('Blends artificial intelligence with human expertise', style='List Bullet 2')
doc.add_paragraph('First AI-native contact center platform', style='List Bullet 2')
doc.add_paragraph('Provides both technology platform AND outsourced agents', style='List Bullet 2')

doc.add_heading('Market Position', 2)
doc.add_paragraph('Valuation: $500 million (as of October 2024)', style='List Bullet')
doc.add_paragraph('Total Funding: $50 million in total financing', style='List Bullet')
doc.add_paragraph('Lead Investor: General Catalyst (Series C)', style='List Bullet')
doc.add_paragraph('ARR Target: Exceeding $100M ARR (announced September 2025)', style='List Bullet')
doc.add_paragraph('Global Reach: Operations across 6 continents', style='List Bullet')
doc.add_paragraph('Customer Base: 200+ customers (post-PartnerHero acquisition)', style='List Bullet')

doc.add_heading('Total Addressable Market', 2)
p = doc.add_paragraph('Disrupting the ')
run = p.add_run('$741 billion global contact center market')
run.bold = True

doc.add_page_break()

# Recent News & Milestones
doc.add_heading('Recent News & Milestones', 1)

doc.add_heading('Major Funding (October 2024)', 2)
doc.add_paragraph('$50M Series C led by General Catalyst at $500M valuation', style='List Bullet')
doc.add_paragraph('Additional investors: Celesta Capital, Alorica, and angel investors', style='List Bullet')
doc.add_paragraph('Less than one year old when securing this funding', style='List Bullet')

doc.add_heading('Strategic Acquisition', 2)
doc.add_paragraph('Acquired PartnerHero: Leading customer operations outsourcing provider', style='List Bullet')
doc.add_paragraph('Added 200+ customers and staff worldwide', style='List Bullet')
doc.add_paragraph('Enabled rapid global expansion to 6 continents', style='List Bullet')

doc.add_heading('Growth Trajectory (2025)', 2)
doc.add_paragraph('Plans to exceed $100M ARR', style='List Bullet')
doc.add_paragraph('Accelerating global adoption of AI-native contact center', style='List Bullet')
doc.add_paragraph('Positioned as industry disruptor', style='List Bullet')

doc.add_heading('Company Origins', 2)
doc.add_paragraph('Originally conceived by President and co-founder Anand Chandrasekaran', style='List Bullet')
doc.add_paragraph('Anand was a Partner at General Catalyst before founding Crescendo', style='List Bullet')

doc.add_page_break()

# Competitive Landscape
doc.add_heading('Competitive Landscape', 1)

doc.add_heading('Direct AI CX Platform Competitors', 2)

p = doc.add_paragraph()
p.add_run('Tier 1 - Major Players:').bold = True
doc.add_paragraph('Sierra - AI agents for customer support across business sectors', style='List Bullet')
doc.add_paragraph('Cognigy - Conversational AI platforms and customer service automation', style='List Bullet')
doc.add_paragraph('Decagon - Generative AI for customer support', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Tier 2 - Established Platforms:').bold = True
doc.add_paragraph('Ada - AI-powered customer service automation', style='List Bullet')
doc.add_paragraph('Replicant - Contact Center Automation (CCA)', style='List Bullet')
doc.add_paragraph('Zendesk - Leading provider of customer relationship software', style='List Bullet')
doc.add_paragraph('Zoom Virtual Agent - Conversational AI with NLP', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Tier 3 - Specialized Solutions:').bold = True
doc.add_paragraph('GenieTalk - Advanced conversational AI for contact centers', style='List Bullet')
doc.add_paragraph('Inbenta - AI chatbots, search tools, and knowledge management', style='List Bullet')
doc.add_paragraph('OneReach.ai - No-code conversational AI platforms', style='List Bullet')

doc.add_heading('Crescendo\'s Differentiation', 2)
doc.add_paragraph('Unique "augmented-AI" approach combining AI + human experts', style='List Bullet')
doc.add_paragraph('Pay-for-outcomes business model vs. traditional SaaS pricing', style='List Bullet')
doc.add_paragraph('Total Outcome Guarantee on quality, speed, and satisfaction', style='List Bullet')
doc.add_paragraph('Fully managed solution (technology + people)', style='List Bullet')
doc.add_paragraph('Rapid scaling capability (6 continents in <1 year)', style='List Bullet')

doc.add_page_break()

# Enterprise Account Executive Role
doc.add_heading('Enterprise Account Executive Role', 1)

doc.add_heading('Position Overview', 2)
doc.add_paragraph('High-impact role leading relationships with large, complex enterprises and guiding them through AI transformation of their customer experience operations.')

doc.add_heading('Key Responsibilities', 2)

p = doc.add_paragraph()
p.add_run('1. Drive Sales Excellence').bold = True
doc.add_paragraph('Develop and execute strategic plans to uncover new opportunities', style='List Bullet')
doc.add_paragraph('Build relationships across multiple departments (CX, IT, executive leadership)', style='List Bullet')
doc.add_paragraph('Navigate complex enterprise buying processes', style='List Bullet')

p = doc.add_paragraph()
p.add_run('2. Solution Selling').bold = True
doc.add_paragraph('Position Crescendo\'s Augmented AI solutions for enterprise CX challenges', style='List Bullet')
doc.add_paragraph('Tailor presentations, demos, and proposals to unique client needs', style='List Bullet')
doc.add_paragraph('Clearly articulate ROI and business value', style='List Bullet')
doc.add_paragraph('Manage multi-million dollar deal cycles', style='List Bullet')

p = doc.add_paragraph()
p.add_run('3. Customer Engagement').bold = True
doc.add_paragraph('Build long-term relationships with C-level executives', style='List Bullet')
doc.add_paragraph('Serve as trusted advisor on industry trends and best practices', style='List Bullet')
doc.add_paragraph('Provide insights on customer experience, AI, and automation', style='List Bullet')

p = doc.add_paragraph()
p.add_run('4. Cross-Functional Collaboration').bold = True
doc.add_paragraph('Work with product, marketing, and customer success teams', style='List Bullet')
doc.add_paragraph('Ensure alignment on client needs and expectations', style='List Bullet')
doc.add_paragraph('Facilitate seamless delivery of solutions post-sale', style='List Bullet')

p = doc.add_paragraph()
p.add_run('5. Market Expansion').bold = True
doc.add_paragraph('Identify and pursue new business opportunities in assigned region', style='List Bullet')
doc.add_paragraph('Expand Crescendo\'s footprint within key accounts and verticals', style='List Bullet')
doc.add_paragraph('Drive land-and-expand strategies', style='List Bullet')

p = doc.add_paragraph()
p.add_run('6. Thought Leadership').bold = True
doc.add_paragraph('Represent Crescendo at industry events and conferences', style='List Bullet')
doc.add_paragraph('Share insights, case studies, and success stories', style='List Bullet')
doc.add_paragraph('Become recognized expert in AI and customer experience space', style='List Bullet')

doc.add_page_break()

doc.add_heading('Required Qualifications', 2)

p = doc.add_paragraph()
p.add_run('Experience:').bold = True
doc.add_paragraph('3+ years in Enterprise sales (CX, AI, or similar technology)', style='List Bullet')
doc.add_paragraph('Strong history of closing complex, multi-million-dollar deals', style='List Bullet')
doc.add_paragraph('Experience managing relationships with large, global enterprises', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Sales Skills:').bold = True
doc.add_paragraph('Consultative selling approach', style='List Bullet')
doc.add_paragraph('Deep understanding of sales process and pipeline management', style='List Bullet')
doc.add_paragraph('Proven ability to engage C-level decision-makers', style='List Bullet')
doc.add_paragraph('Strong negotiation and closing skills', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Personal Attributes:').bold = True
doc.add_paragraph('Self-starter with ability to take ownership and lead independently', style='List Bullet')
doc.add_paragraph('Exceptional interpersonal and communication skills', style='List Bullet')
doc.add_paragraph('Strong relationship-building capabilities', style='List Bullet')
doc.add_paragraph('Critical and creative thinking for solution design', style='List Bullet')
doc.add_paragraph('Adaptable to fast-paced, dynamic environment', style='List Bullet')

p = doc.add_paragraph()
p.add_run('Work Arrangement:').bold = True
doc.add_paragraph('Hybrid position with location flexibility', style='List Bullet')
doc.add_paragraph('Proximity to major markets is a plus', style='List Bullet')
doc.add_paragraph('Travel required for client meetings and key events', style='List Bullet')

doc.add_page_break()

# MEDDPICC
doc.add_heading('MEDDPICC/STAR Talking Points', 1)

doc.add_heading('Metrics (M)', 2)
doc.add_paragraph('Your quantifiable achievements to highlight:')
doc.add_paragraph('Revenue growth you\'ve driven (YoY %, quota attainment %)', style='List Bullet')
doc.add_paragraph('Deal sizes you\'ve closed ($XM ACV/TCV)', style='List Bullet')
doc.add_paragraph('Pipeline generation and conversion rates', style='List Bullet')
doc.add_paragraph('Customer retention/expansion metrics', style='List Bullet')
doc.add_paragraph('Sales cycle reduction achievements', style='List Bullet')

doc.add_heading('Economic Buyer (E)', 2)
doc.add_paragraph('At Crescendo, likely decision-makers include:')
doc.add_paragraph('Chief Customer Officer (CCO)', style='List Bullet')
doc.add_paragraph('Chief Experience Officer (CXO)', style='List Bullet')
doc.add_paragraph('VP of Customer Support/Success', style='List Bullet')
doc.add_paragraph('CTO/CIO (technology evaluation)', style='List Bullet')
doc.add_paragraph('CFO (ROI approval for multi-million dollar deals)', style='List Bullet')

doc.add_heading('Decision Criteria (D)', 2)
doc.add_paragraph('What matters to Crescendo\'s prospects:')
doc.add_paragraph('ROI and cost savings vs. traditional contact centers', style='List Bullet')
doc.add_paragraph('Quality of customer experience outcomes', style='List Bullet')
doc.add_paragraph('Speed of implementation and time-to-value', style='List Bullet')
doc.add_paragraph('Scalability and multilingual capabilities', style='List Bullet')
doc.add_paragraph('AI accuracy and human backup reliability', style='List Bullet')
doc.add_paragraph('Integration with existing tech stack', style='List Bullet')

doc.add_heading('Identify Pain (I)', 2)
doc.add_paragraph('Problems Crescendo solves:')
doc.add_paragraph('High cost of traditional contact centers and BPOs', style='List Bullet')
doc.add_paragraph('Poor customer experience from chatbots that can\'t handle complexity', style='List Bullet')
doc.add_paragraph('Inability to scale support globally with quality', style='List Bullet')
doc.add_paragraph('Lack of 24/7 multilingual support', style='List Bullet')
doc.add_paragraph('Difficulty measuring and guaranteeing CX outcomes', style='List Bullet')

doc.add_page_break()

# Strategic Questions
doc.add_heading('Strategic Questions to Ask', 1)

doc.add_heading('About the Role', 2)
doc.add_paragraph('What does success look like for an Enterprise AE in the first 90 days? First year?', style='List Number')
doc.add_paragraph('How is the sales territory structured, and what\'s the target account profile?', style='List Number')
doc.add_paragraph('What\'s the typical deal size and sales cycle for enterprise customers?', style='List Number')
doc.add_paragraph('How does Crescendo support AEs with solution engineering and customer success resources?', style='List Number')
doc.add_paragraph('What are the quota expectations and compensation structure?', style='List Number')

doc.add_heading('About the Company', 2)
doc.add_paragraph('With the recent $50M raise and PartnerHero acquisition, what are Crescendo\'s top 3 priorities for 2026?', style='List Number')
doc.add_paragraph('How is Crescendo differentiating in the market given competitors like Sierra, Cognigy, and Decagon?', style='List Number')
doc.add_paragraph('What verticals or use cases are seeing the strongest adoption of the platform?', style='List Number')
doc.add_paragraph('How does the "pay-for-outcomes" model resonate with enterprise buyers vs. traditional SaaS pricing?', style='List Number')
doc.add_paragraph('What\'s the product roadmap looking like, and how does customer feedback influence it?', style='List Number')

doc.add_heading('About the Team & Culture', 2)
doc.add_paragraph('Can you describe the sales team structure and who I\'d be working most closely with?', style='List Number')
doc.add_paragraph('What\'s the typical background of successful AEs at Crescendo?', style='List Number')
doc.add_paragraph('How does Crescendo maintain its fast-paced startup culture while scaling to 6 continents?', style='List Number')
doc.add_paragraph('What professional development and career growth opportunities exist for top performers?', style='List Number')
doc.add_paragraph('What do you love most about working at Crescendo?', style='List Number')

doc.add_page_break()

# Interview Talking Points
doc.add_heading('Key Talking Points for Interview', 1)

doc.add_heading('Opening Statement (30-60 seconds)', 2)
doc.add_paragraph('"I\'m an enterprise sales professional with [X] years of experience selling complex technology solutions to large enterprises. Most recently at [Company], I focused on [relevant area like CX/AI/SaaS], where I consistently exceeded quota by closing multi-million dollar deals with Fortune 500 companies. What excites me about Crescendo is the unique augmented-AI approach—combining the best of AI automation with human expertise—and the pay-for-outcomes model, which aligns perfectly with how I\'ve always sold: focusing on tangible business value and ROI. I\'m particularly drawn to Crescendo\'s rapid growth trajectory and the opportunity to shape the future of customer experience in a $741B market."')

doc.add_heading('Why Crescendo?', 2)
doc.add_paragraph('Market Leadership: "Crescendo is pioneering the augmented-AI category in CX, and the recent $50M raise at $500M valuation validates the massive opportunity."', style='List Number')
doc.add_paragraph('Differentiated Solution: "The combination of AI + human experts is the right approach—chatbots alone fail, and pure BPOs are too expensive."', style='List Number')
doc.add_paragraph('Proven Growth: "Going from zero to 6 continents and 200+ customers in less than a year shows exceptional execution."', style='List Number')
doc.add_paragraph('Alignment with My Strengths: "I thrive in consultative, complex sales with C-level buyers."', style='List Number')

doc.add_heading('Why You Should Hire Me', 2)
doc.add_paragraph('Track Record: "I have [X years] experience closing enterprise deals averaging $[Y]M, with [Z%] quota attainment."', style='List Number')
doc.add_paragraph('Relevant Experience: "My background in [CX/AI/SaaS] means I understand the buyer, the pain points, and how to articulate ROI."', style='List Number')
doc.add_paragraph('Relationship Builder: "I\'ve built trusted advisor relationships with C-level executives at companies like [examples]."', style='List Number')
doc.add_paragraph('Fast-Growth DNA: "I\'ve thrived at [startup/high-growth company], where I [specific achievement]."', style='List Number')

doc.add_heading('Closing Statement', 2)
doc.add_paragraph('"I\'m genuinely excited about this opportunity. Crescendo is at an inflection point—proven product-market fit, strong funding, and massive market opportunity. I believe my experience in enterprise sales, my consultative approach, and my passion for AI-driven transformation make me the right person to help Crescendo capture this market. I\'m ready to contribute from day one."')

doc.add_page_break()

# Add footer
doc.add_heading('Good luck with your interview!', 2)
p = doc.add_paragraph('This comprehensive research document was generated on January 7, 2026.')
p.runs[0].italic = True

# Save
doc.save('output/Crescendo_AI_Enterprise_AE_Research.docx')
print('✓ Word document created successfully!')
print('✓ File saved: output/Crescendo_AI_Enterprise_AE_Research.docx')
