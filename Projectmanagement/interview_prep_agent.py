#!/usr/bin/env python3
"""
Interview Prep Agent - AI-powered interview research assistant
Uses Claude API with tool calling for comprehensive company and role research
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List
from anthropic import Anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import configuration
from config import (
    ANTHROPIC_API_KEY,
    MODEL_NAME,
    MAX_TOKENS,
    TEMPERATURE,
    OUTPUT_DIR,
    TEMPLATES_DIR
)


class InterviewPrepAgent:
    """Main agent class for interview preparation research"""

    def __init__(self):
        """Initialize the agent with Anthropic client and tools"""
        self.client = Anthropic(api_key=ANTHROPIC_API_KEY)
        self.conversation_history = []
        self.research_data = {}

        # Ensure output directory exists
        Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
        Path(TEMPLATES_DIR).mkdir(parents=True, exist_ok=True)

        # Define available tools
        self.tools = [
            {
                "name": "web_search",
                "description": "Search the web for information about companies, roles, news, and competitors. Returns relevant search results.",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "query": {
                            "type": "string",
                            "description": "The search query to execute"
                        },
                        "num_results": {
                            "type": "integer",
                            "description": "Number of results to return (default: 5)",
                            "default": 5
                        }
                    },
                    "required": ["query"]
                }
            },
            {
                "name": "save_research",
                "description": "Save research findings to the internal database for later retrieval and document generation",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "category": {
                            "type": "string",
                            "description": "Research category (e.g., 'company_overview', 'role_requirements', 'competitors', 'recent_news')"
                        },
                        "content": {
                            "type": "string",
                            "description": "The research content to save"
                        },
                        "source": {
                            "type": "string",
                            "description": "Source URL or reference"
                        }
                    },
                    "required": ["category", "content"]
                }
            },
            {
                "name": "generate_word_document",
                "description": "Generate a professional Word document with the research findings",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "company_name": {
                            "type": "string",
                            "description": "Name of the company"
                        },
                        "role_title": {
                            "type": "string",
                            "description": "Job title/role"
                        },
                        "filename": {
                            "type": "string",
                            "description": "Optional custom filename (without extension)"
                        }
                    },
                    "required": ["company_name", "role_title"]
                }
            },
            {
                "name": "generate_markdown",
                "description": "Generate a markdown file with research findings for easy editing",
                "input_schema": {
                    "type": "object",
                    "properties": {
                        "company_name": {
                            "type": "string",
                            "description": "Name of the company"
                        },
                        "role_title": {
                            "type": "string",
                            "description": "Job title/role"
                        },
                        "filename": {
                            "type": "string",
                            "description": "Optional custom filename (without extension)"
                        }
                    },
                    "required": ["company_name", "role_title"]
                }
            }
        ]

    def execute_tool(self, tool_name: str, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Execute a tool and return results"""

        if tool_name == "web_search":
            return self._mock_web_search(tool_input)

        elif tool_name == "save_research":
            return self._save_research(tool_input)

        elif tool_name == "generate_word_document":
            return self._generate_word_document(tool_input)

        elif tool_name == "generate_markdown":
            return self._generate_markdown(tool_input)

        else:
            return {"error": f"Unknown tool: {tool_name}"}

    def _mock_web_search(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Mock web search - in production, integrate with real search API"""
        query = tool_input.get("query", "")
        num_results = tool_input.get("num_results", 5)

        # This is a mock implementation
        # In production, integrate with services like:
        # - Exa AI (exa.ai)
        # - SerpAPI
        # - Google Custom Search API
        # - Bing Search API

        return {
            "query": query,
            "results": [
                {
                    "title": f"Result for: {query}",
                    "url": "https://example.com",
                    "snippet": "Mock search result. Integrate with real search API for production use."
                }
            ],
            "note": "This is a mock search. Please integrate with a real search API (Exa AI, SerpAPI, etc.) for production use."
        }

    def _save_research(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Save research findings to internal database"""
        category = tool_input.get("category")
        content = tool_input.get("content")
        source = tool_input.get("source", "")

        if category not in self.research_data:
            self.research_data[category] = []

        self.research_data[category].append({
            "content": content,
            "source": source,
            "timestamp": datetime.now().isoformat()
        })

        return {
            "status": "success",
            "message": f"Saved to {category}",
            "total_items": len(self.research_data[category])
        }

    def _generate_word_document(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Generate a professional Word document with research findings"""
        company_name = tool_input.get("company_name")
        role_title = tool_input.get("role_title")
        custom_filename = tool_input.get("filename")

        # Create filename
        if custom_filename:
            filename = f"{custom_filename}.docx"
        else:
            safe_company = company_name.replace(" ", "_")
            safe_role = role_title.replace(" ", "_")
            timestamp = datetime.now().strftime("%Y%m%d")
            filename = f"{safe_company}_{safe_role}_Research_{timestamp}.docx"

        filepath = Path(OUTPUT_DIR) / filename

        # Create document
        doc = Document()

        # Add title
        title = doc.add_heading(f"{company_name} - {role_title}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add subtitle with date
        subtitle = doc.add_paragraph(f"Interview Preparation Research")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Add research sections
        for category, items in self.research_data.items():
            # Format category name
            section_title = category.replace("_", " ").title()
            doc.add_heading(section_title, 1)

            for item in items:
                # Add content
                p = doc.add_paragraph(item['content'])

                # Add source if available
                if item.get('source'):
                    source_p = doc.add_paragraph(f"Source: {item['source']}")
                    source_p.runs[0].font.size = Pt(9)
                    source_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

                doc.add_paragraph()  # Spacer

        # Save document
        doc.save(str(filepath))

        return {
            "status": "success",
            "filepath": str(filepath),
            "message": f"Word document generated: {filename}"
        }

    def _generate_markdown(self, tool_input: Dict[str, Any]) -> Dict[str, Any]:
        """Generate markdown file with research findings"""
        company_name = tool_input.get("company_name")
        role_title = tool_input.get("role_title")
        custom_filename = tool_input.get("filename")

        # Create filename
        if custom_filename:
            filename = f"{custom_filename}.md"
        else:
            safe_company = company_name.replace(" ", "_")
            safe_role = role_title.replace(" ", "_")
            timestamp = datetime.now().strftime("%Y%m%d")
            filename = f"{safe_company}_{safe_role}_Research_{timestamp}.md"

        filepath = Path(OUTPUT_DIR) / filename

        # Build markdown content
        md_content = []
        md_content.append(f"# {company_name} - {role_title}\n")
        md_content.append(f"**Interview Preparation Research**\n")
        md_content.append(f"*Generated: {datetime.now().strftime('%B %d, %Y')}*\n")
        md_content.append("\n---\n\n")

        # Add research sections
        for category, items in self.research_data.items():
            section_title = category.replace("_", " ").title()
            md_content.append(f"## {section_title}\n\n")

            for item in items:
                md_content.append(f"{item['content']}\n\n")

                if item.get('source'):
                    md_content.append(f"*Source: {item['source']}*\n\n")

                md_content.append("---\n\n")

        # Write file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.writelines(md_content)

        return {
            "status": "success",
            "filepath": str(filepath),
            "message": f"Markdown file generated: {filename}"
        }

    def chat(self, user_message: str) -> str:
        """Process user message and return agent response"""

        # Add user message to history
        self.conversation_history.append({
            "role": "user",
            "content": user_message
        })

        # System prompt
        system_prompt = """You are an expert interview preparation assistant. Your role is to help users research companies and prepare for job interviews.

When a user asks you to research a company or role, you should:
1. Use web_search to find comprehensive information about the company, role, recent news, and competitors
2. Save important findings using save_research with appropriate categories
3. Generate professional documents when requested

Research categories to use:
- company_overview: Company mission, values, products, culture
- role_requirements: Job description, required skills, responsibilities
- recent_news: Latest company news, funding, product launches
- competitors: Competitive landscape and differentiation
- talking_points: MEDDPICC/STAR-aligned interview talking points
- questions: Thoughtful questions to ask the interviewer

Be thorough, proactive, and conversational. Anticipate what information would be most valuable for interview preparation."""

        # Make API call with tools
        response = self.client.messages.create(
            model=MODEL_NAME,
            max_tokens=MAX_TOKENS,
            temperature=TEMPERATURE,
            system=system_prompt,
            messages=self.conversation_history,
            tools=self.tools
        )

        # Process response and handle tool calls
        while response.stop_reason == "tool_use":
            # Extract assistant message and tool uses
            assistant_content = []
            tool_results = []

            for block in response.content:
                if block.type == "text":
                    assistant_content.append(block)
                elif block.type == "tool_use":
                    # Execute tool
                    tool_result = self.execute_tool(block.name, block.input)

                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": json.dumps(tool_result)
                    })

                    assistant_content.append(block)

            # Add assistant message to history
            self.conversation_history.append({
                "role": "assistant",
                "content": assistant_content
            })

            # Add tool results to history
            self.conversation_history.append({
                "role": "user",
                "content": tool_results
            })

            # Continue conversation
            response = self.client.messages.create(
                model=MODEL_NAME,
                max_tokens=MAX_TOKENS,
                temperature=TEMPERATURE,
                system=system_prompt,
                messages=self.conversation_history,
                tools=self.tools
            )

        # Extract final response text
        final_response = ""
        for block in response.content:
            if block.type == "text":
                final_response += block.text

        # Add final assistant message to history
        self.conversation_history.append({
            "role": "assistant",
            "content": response.content
        })

        return final_response


def main():
    """Main CLI interface"""
    print("=" * 60)
    print("🎯 Interview Prep Agent")
    print("=" * 60)
    print("\nAI-powered interview research assistant")
    print("Type 'quit' or 'exit' to end the session\n")

    # Initialize agent
    agent = InterviewPrepAgent()

    # Main conversation loop
    while True:
        try:
            # Get user input
            user_input = input(">> ").strip()

            # Check for exit
            if user_input.lower() in ['quit', 'exit', 'q']:
                print("\n👋 Goodbye! Good luck with your interviews!\n")
                break

            # Skip empty input
            if not user_input:
                continue

            # Process message
            print("\n🤔 Thinking...\n")
            response = agent.chat(user_input)

            # Display response
            print(response)
            print("\n" + "-" * 60 + "\n")

        except KeyboardInterrupt:
            print("\n\n👋 Goodbye! Good luck with your interviews!\n")
            break
        except Exception as e:
            print(f"\n❌ Error: {str(e)}\n")
            print("Please try again or type 'quit' to exit.\n")


if __name__ == "__main__":
    main()
